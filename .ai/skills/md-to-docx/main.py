"""
md-to-docx — Convert Markdown files into Word .docx files.

Usage:
    python main.py <input.md> [output.docx]
"""

import sys
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown import markdown

PALETTE = {
    "surface": "FFFFFF",
    "surface_2": "F1F4F8",
    "text": "1F2937",
    "muted": "5B6574",
    "border": "D9E0E7",
    "accent": "0F766E",
    "accent_soft": "E6F4F1",
    "warning": "9A3412",
    "title": "1E3A5F",
    "th_bg": "E6F4F8",
    "table_even_row": "F8FAFC",
    "code_bg": "0F172A",
    "code_text": "E5EEFC",
}


class Converter:
    """Convert markdown documents into .docx files."""

    def convert(self, md_path: Path, out_path: Path) -> Path:
        markdown_text = md_path.read_text(encoding="utf-8")
        html = markdown(
            markdown_text,
            extensions=["tables", "fenced_code", "sane_lists", "nl2br"],
        )
        soup = BeautifulSoup(html, "html.parser")

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        self._apply_document_defaults(doc)

        for node in soup.contents:
            self._add_block(doc, node)

        doc.save(out_path)
        return out_path

    def _apply_document_defaults(self, doc: Document) -> None:
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Segoe UI"
        normal_style.font.size = Pt(11)
        normal_style.font.color.rgb = RGBColor.from_string(PALETTE["text"])
        normal_style.paragraph_format.space_after = Pt(6)

        if "Cell List Bullet" not in [style.name for style in doc.styles]:
            style = doc.styles.add_style("Cell List Bullet", WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles["Normal"]
            style.paragraph_format.left_indent = Pt(0)
            style.paragraph_format.first_line_indent = Pt(0)
            style.paragraph_format.space_after = Pt(6)
            style.font.name = "Segoe UI"
            style.font.size = Pt(10)
            style.font.color.rgb = RGBColor.from_string(PALETTE["text"])

    def _set_run_color(self, run, color_key: str, *, bold: bool = False) -> None:
        run.font.color.rgb = RGBColor.from_string(PALETTE[color_key])
        run.bold = bold

    def _set_paragraph_spacing(
        self, paragraph, *, before: int = 0, after: int = 0
    ) -> None:
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)

    def _set_cell_shading(self, cell, color_hex: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color_hex)
        tc_pr.append(shd)

    def _set_row_height(self, row, height_twips: int) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        tr_height = OxmlElement("w:trHeight")
        tr_height.set(qn("w:val"), str(height_twips))
        tr_height.set(qn("w:hRule"), "atLeast")
        tr_pr.append(tr_height)

    def _set_cell_vertical_alignment(self, cell) -> None:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _set_cell_text(
        self, cell, text: str, *, color_key: str = "text", bold: bool = False
    ) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(text.strip())
        run.font.name = "Segoe UI"
        run.font.size = Pt(10)
        self._set_run_color(run, color_key, bold=bold)

    def _extract_text(self, node: Tag) -> str:
        return " ".join(node.stripped_strings)

    def _extract_lines(self, node: Tag) -> list[str]:
        html = node.decode_contents()
        html = html.replace("<br/>", "\n").replace("<br />", "\n").replace("<br>", "\n")
        text = BeautifulSoup(html, "html.parser").get_text("\n")
        return [line.strip() for line in text.splitlines() if line.strip()]

    def _extract_cell_lines(self, node: Tag) -> list[str]:
        return self._extract_lines(node)

    def _append_inline_content(self, paragraph, node) -> None:
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                run = paragraph.add_run(text)
                run.font.name = "Segoe UI"
                run.font.size = Pt(11)
                self._set_run_color(run, "text")
            return

        if not isinstance(node, Tag):
            return

        name = node.name.lower()
        if name == "br":
            paragraph.add_run().add_break(WD_BREAK.LINE)
            return

        if name == "strong":
            run = paragraph.add_run(node.get_text())
            run.font.name = "Segoe UI"
            run.font.size = Pt(11)
            self._set_run_color(run, "text", bold=True)
            return

        if name == "a":
            run = paragraph.add_run(node.get_text())
            run.font.name = "Segoe UI"
            run.font.size = Pt(11)
            self._set_run_color(run, "accent")
            run.underline = True
            return

        for child in node.children:
            self._append_inline_content(paragraph, child)

    def _populate_cell(self, cell, node: Tag, *, is_header: bool) -> None:
        cell.text = ""
        lines = self._extract_cell_lines(node)
        if not lines:
            return

        first_paragraph = True
        for line in lines:
            is_bullet = line.startswith("- ")
            text = line[2:].strip() if is_bullet else line

            paragraph = cell.paragraphs[0] if first_paragraph else cell.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            first_paragraph = False

            if is_header:
                self._set_cell_vertical_alignment(cell)
                run = paragraph.add_run(text)
                run.font.name = "Segoe UI"
                run.font.size = Pt(10)
                self._set_run_color(run, "title", bold=True)
                continue

            if is_bullet:
                paragraph.style = "Cell List Bullet"
                run = paragraph.add_run(f"• {text}")
            else:
                run = paragraph.add_run(text)

            run.font.name = "Segoe UI"
            run.font.size = Pt(10)
            self._set_run_color(run, "text")

    def _add_list(self, doc: Document, node: Tag, style: str) -> None:
        for li in node.find_all("li", recursive=False):
            paragraph = doc.add_paragraph(self._extract_text(li), style=style)
            for run in paragraph.runs:
                run.font.name = "Segoe UI"
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor.from_string(PALETTE["text"])

    def _set_table_column_widths(self, table, col_count: int) -> None:
        if col_count == 2:
            widths = [0.34, 0.66]
        elif col_count == 3:
            widths = [0.20, 0.40, 0.40]
        elif col_count == 4:
            widths = [0.25, 0.25, 0.25, 0.25]
        else:
            return

        section_width = 6.9
        for col_index, ratio in enumerate(widths):
            width = Inches(section_width * ratio)
            for cell in table.columns[col_index].cells:
                cell.width = width

    def _add_table(self, doc: Document, node: Tag) -> None:
        rows = []
        for tr in node.find_all("tr"):
            cells = tr.find_all(["th", "td"], recursive=False)
            rows.append(cells)
        if not rows:
            return

        col_count = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_column_widths(table, col_count)

        if table.rows:
            self._set_row_height(table.rows[0], 600)

        for row_index, row in enumerate(rows):
            for col_index in range(col_count):
                cell = table.cell(row_index, col_index)
                source_node = row[col_index] if col_index < len(row) else None
                is_header = row_index == 0
                if is_header:
                    self._set_cell_shading(cell, PALETTE["th_bg"])
                elif row_index % 2 == 0:
                    self._set_cell_shading(cell, PALETTE["table_even_row"])

                if source_node is not None:
                    self._populate_cell(cell, source_node, is_header=is_header)

    def _add_block(self, doc: Document, node) -> None:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                doc.add_paragraph(text)
            return

        if not isinstance(node, Tag):
            return

        name = node.name.lower()
        if name == "h1":
            paragraph = doc.add_heading(level=0)
            self._set_paragraph_spacing(paragraph, after=10)
            run = paragraph.add_run(self._extract_text(node))
            run.font.name = "Segoe UI"
            run.font.size = Pt(24)
            self._set_run_color(run, "accent", bold=True)
        elif name == "h2":
            paragraph = doc.add_heading(level=1)
            self._set_paragraph_spacing(paragraph, after=8)
            run = paragraph.add_run(self._extract_text(node))
            run.font.name = "Segoe UI"
            run.font.size = Pt(18)
            self._set_run_color(run, "title", bold=True)
        elif name == "h3":
            paragraph = doc.add_heading(level=2)
            self._set_paragraph_spacing(paragraph, after=6)
            run = paragraph.add_run(self._extract_text(node))
            run.font.name = "Segoe UI"
            run.font.size = Pt(14)
            self._set_run_color(run, "title", bold=True)
        elif name == "h4":
            paragraph = doc.add_heading(level=3)
            run = paragraph.add_run(self._extract_text(node))
            run.font.name = "Segoe UI"
            run.font.size = Pt(12)
            self._set_run_color(run, "title", bold=True)
        elif name == "p":
            paragraph = doc.add_paragraph()
            self._set_paragraph_spacing(paragraph, after=6)
            for child in node.children:
                self._append_inline_content(paragraph, child)
        elif name == "ul":
            self._add_list(doc, node, "List Bullet")
        elif name == "ol":
            self._add_list(doc, node, "List Number")
        elif name == "table":
            self._add_table(doc, node)
        elif name == "hr":
            doc.add_paragraph("")
        elif name == "blockquote":
            text = self._extract_text(node)
            if text:
                paragraph = doc.add_paragraph()
                paragraph.style = (
                    "Intense Quote"
                    if "Intense Quote" in [style.name for style in doc.styles]
                    else "Quote"
                )
                run = paragraph.add_run(text)
                run.font.name = "Segoe UI"
                run.font.size = Pt(11)
                self._set_run_color(run, "muted")
        elif name == "pre":
            text = node.get_text("", strip=False).rstrip()
            if text:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(text)
                run.font.name = "Cascadia Code"
                run.font.size = Pt(10)
                self._set_run_color(run, "code_text")
        else:
            for child in node.children:
                self._add_block(doc, child)


def main() -> None:
    if len(sys.argv) < 2:
        print("Usage: python main.py <input.md> [output.docx]")
        sys.exit(1)

    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"Error: file not found: {md_path}")
        sys.exit(1)

    out_path = Path(sys.argv[2]) if len(sys.argv) >= 3 else md_path.with_suffix(".docx")
    result = Converter().convert(md_path, out_path)
    print(f"DOCX written to {result} ({result.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
