"""
md-to-docx — Convert Markdown files into Word .docx files.

Usage:
    python main.py <input.md> [output.docx]
"""

import sys
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown import markdown

PALETTE = {
    "surface": "FFFFFF",
    "surface_2": "F1F4F8",
    "text": "1F2937",
    "muted": "5B6574",
    "border": "D9E0E7",
    "accent": "0F766E",
    "accent_soft": "E6F4F1",
    "warning": "9A3412",
    "title": "1E3A5F",
    "th_bg": "E6F4F8",
    "table_even_row": "F8FAFC",
    "code_bg": "0F172A",
    "code_text": "E5EEFC",
}


class Converter:
    """Convert markdown documents into .docx files."""

    def convert(self, md_path: Path, out_path: Path) -> Path:
        """Convert a Markdown file to .docx. Returns the output path."""
        try:
            markdown_text = md_path.read_text(encoding="utf-8")
        except OSError as exc:
            raise OSError(f"Failed to read input file: {md_path}") from exc

        try:
            html = markdown(
                markdown_text,
                extensions=["tables", "fenced_code", "sane_lists", "nl2br"],
            )
        except Exception as exc:
            raise ValueError(f"Failed to parse Markdown in {md_path}: {exc}") from exc

        soup = BeautifulSoup(html, "lxml")

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        self._apply_document_defaults(doc)

        for node in soup.contents:
            self._add_block(doc, node)

        doc.save(out_path)
        return out_path

    def _apply_document_defaults(self, doc: Document) -> None:
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Segoe UI"
        normal_style.font.size = Pt(11)
        normal_style.font.color.rgb = RGBColor.from_string(PALETTE["text"])
        normal_style.paragraph_format.space_after = Pt(6)

        if "Cell List Bullet" not in [style.name for style in doc.styles]:
            style = doc.styles.add_style("Cell List Bullet", WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles["Normal"]
            style.paragraph_format.left_indent = Pt(0)
            style.paragraph_format.first_line_indent = Pt(0)
            style.paragraph_format.space_after = Pt(6)
            style.font.name = "Segoe UI"
            style.font.size = Pt(10)
            style.font.color.rgb = RGBColor.from_string(PALETTE["text"])

    def _set_run_color(self, run, color_key: str, *, bold: bool = False) -> None:
        run.font.color.rgb = RGBColor.from_string(PALETTE[color_key])
        run.bold = bold

    def _add_run(
        self,
        paragraph,
        text: str,
        *,
        font_name: str = "Segoe UI",
        font_size: int = 11,
        color: str = "text",
        bold: bool = False,
    ):
        """Add a formatted run to the paragraph. Returns the run for further customization."""
        run = paragraph.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        self._set_run_color(run, color, bold=bold)
        return run

    def _set_paragraph_spacing(
        self, paragraph, *, before: int = 0, after: int = 0
    ) -> None:
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)

    def _set_cell_shading(self, cell, color_hex: str) -> None:
        """Set cell background color. Uses private API — python-docx has no public shading API (tested: 1.2.0)."""
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color_hex)
        tc_pr.append(shd)

    def _set_row_height(self, row, height_twips: int) -> None:
        """Set minimum row height. Uses private API — python-docx has no public row height API (tested: 1.2.0)."""
        tr_pr = row._tr.get_or_add_trPr()
        tr_height = OxmlElement("w:trHeight")
        tr_height.set(qn("w:val"), str(height_twips))
        tr_height.set(qn("w:hRule"), "atLeast")
        tr_pr.append(tr_height)

    def _set_cell_vertical_alignment(self, cell) -> None:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _add_hyperlink(
        self, paragraph, url: str, text: str, font_size: int = 11
    ) -> None:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        c = OxmlElement("w:color")
        c.set(qn("w:val"), PALETTE["accent"])
        rPr.append(c)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Segoe UI")
        rFonts.set(qn("w:hAnsi"), "Segoe UI")
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(font_size * 2))
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        hyperlink.append(r)
        # Uses private API — python-docx has no public hyperlink API (tested: 1.2.0)
        paragraph._p.append(hyperlink)

    def _extract_text(self, node: Tag) -> str:
        return " ".join(node.stripped_strings)

    def _extract_lines(self, node: Tag) -> list[str]:
        text = node.get_text("\n")
        return [line.strip() for line in text.splitlines() if line.strip()]

    def _append_inline_content(self, paragraph, node) -> None:
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                # Skip newline-only strings between inline elements inside <p>
                stripped = text.replace("\n", "").replace("\r", "")
                if stripped:
                    self._add_run(paragraph, text)
                # else: pure newline between tags — skip, don't add anything
            return

        if not isinstance(node, Tag):
            return

        name = node.name.lower()
        if name == "br":
            paragraph.add_run().add_break(WD_BREAK.LINE)
            return

        if name == "strong":
            self._add_run(paragraph, node.get_text(), bold=True)
            return

        if name == "a":
            href = node.get("href", "")
            text = node.get_text()
            if href:
                self._add_hyperlink(paragraph, href, text)
            else:
                run = self._add_run(paragraph, text, color="accent")
                run.underline = True
            return

        children = list(node.children)
        for child in children:
            self._append_inline_content(paragraph, child)

    def _populate_cell(self, cell, node: Tag, *, is_header: bool) -> None:
        cell.text = ""
        lines = self._extract_lines(node)
        if not lines:
            return

        first_paragraph = True
        for line in lines:
            is_bullet = line.startswith("- ")
            text = line[2:].strip() if is_bullet else line

            paragraph = cell.paragraphs[0] if first_paragraph else cell.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            first_paragraph = False

            if is_header:
                self._set_cell_vertical_alignment(cell)
                self._add_run(paragraph, text, font_size=10, color="title", bold=True)
                continue

            if is_bullet:
                paragraph.style = "Cell List Bullet"
                self._add_run(paragraph, f"• {text}", font_size=10)
            else:
                self._add_run(paragraph, text, font_size=10)

    def _add_list(self, doc: Document, node: Tag, style: str) -> None:
        for li in node.find_all("li", recursive=False):
            self._add_list_item(doc, li, style)

    def _add_list_item(self, doc: Document, li: Tag, style: str) -> None:
        paragraph = doc.add_paragraph(style=style)
        self._set_paragraph_spacing(paragraph, after=6)
        for child in li.children:
            if isinstance(child, NavigableString) and not str(child).strip():
                continue
            if isinstance(child, Tag) and child.name in ("ul", "ol"):
                # Nested list — process as separate list items
                nested_style = "List Bullet 2" if style == "List Bullet" else style
                if nested_style not in [s.name for s in doc.styles]:
                    nested_style = style
                self._add_list(doc, child, nested_style)
                continue
            self._append_inline_content(paragraph, child)

    def _set_table_column_widths(self, doc: Document, table, col_count: int) -> None:
        if col_count == 2:
            widths = [0.34, 0.66]
        elif col_count == 3:
            widths = [0.20, 0.40, 0.40]
        elif col_count == 4:
            widths = [0.25, 0.25, 0.25, 0.25]
        else:
            return

        section = doc.sections[0]
        page_width = section.page_width.inches
        left = section.left_margin.inches
        right = section.right_margin.inches
        usable_width = page_width - left - right

        for col_index, ratio in enumerate(widths):
            width = Inches(usable_width * ratio)
            for cell in table.columns[col_index].cells:
                cell.width = width

    def _add_table(self, doc: Document, node: Tag) -> None:
        rows = []
        for tr in node.find_all("tr"):
            cells = tr.find_all(["th", "td"], recursive=False)
            rows.append(cells)
        if not rows:
            return

        col_count = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_column_widths(doc, table, col_count)

        if table.rows:
            self._set_row_height(table.rows[0], 600)

        for row_index, row in enumerate(rows):
            for col_index in range(col_count):
                cell = table.cell(row_index, col_index)
                source_node = row[col_index] if col_index < len(row) else None
                is_header = row_index == 0
                if is_header:
                    self._set_cell_shading(cell, PALETTE["th_bg"])
                elif row_index % 2 == 0:
                    self._set_cell_shading(cell, PALETTE["table_even_row"])

                if source_node is not None:
                    self._populate_cell(cell, source_node, is_header=is_header)

    def _add_heading(
        self,
        doc: Document,
        node: Tag,
        level: int,
        after: int,
        size: int,
        color_key: str,
    ) -> None:
        paragraph = doc.add_heading(level=level)
        if after:
            self._set_paragraph_spacing(paragraph, after=after)
        self._add_run(
            paragraph,
            self._extract_text(node),
            font_size=size,
            color=color_key,
            bold=True,
        )

    def _add_block(self, doc: Document, node) -> None:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                doc.add_paragraph(text)
            return

        if not isinstance(node, Tag):
            return

        name = node.name.lower()
        handler = self._block_handlers.get(name)
        if handler:
            handler(doc, node)
            return

        # Unknown tag — recurse into children
        for child in list(node.children):
            self._add_block(doc, child)

    @property
    def _block_handlers(self):
        return {
            "h1": lambda doc, node: self._add_heading(
                doc, node, level=0, after=10, size=24, color_key="accent"
            ),
            "h2": lambda doc, node: self._add_heading(
                doc, node, level=1, after=8, size=18, color_key="title"
            ),
            "h3": lambda doc, node: self._add_heading(
                doc, node, level=2, after=6, size=14, color_key="title"
            ),
            "h4": lambda doc, node: self._add_heading(
                doc, node, level=3, after=0, size=12, color_key="title"
            ),
            "p": self._add_paragraph_block,
            "ul": lambda doc, node: self._add_list(doc, node, "List Bullet"),
            "ol": lambda doc, node: self._add_list(doc, node, "List Number"),
            "table": self._add_table,
            "hr": lambda doc, _: doc.add_paragraph(""),
            "blockquote": self._add_blockquote_block,
            "pre": self._add_code_block,
        }

    def _add_paragraph_block(self, doc: Document, node: Tag) -> None:
        paragraph = doc.add_paragraph()
        self._set_paragraph_spacing(paragraph, after=6)
        for child in list(node.children):
            self._append_inline_content(paragraph, child)

    def _add_blockquote_block(self, doc: Document, node: Tag) -> None:
        text = self._extract_text(node)
        if text:
            paragraph = doc.add_paragraph()
            paragraph.style = (
                "Intense Quote"
                if "Intense Quote" in [style.name for style in doc.styles]
                else "Quote"
            )
            self._add_run(paragraph, text, color="muted")

    def _add_code_block(self, doc: Document, node: Tag) -> None:
        text = node.get_text("", strip=False).rstrip()
        if text:
            paragraph = doc.add_paragraph()
            self._add_run(
                paragraph,
                text,
                font_name="Cascadia Code",
                font_size=10,
                color="code_text",
            )


def main() -> None:
    if len(sys.argv) < 2:
        print("Usage: python main.py <input.md> [output.docx]")
        sys.exit(1)

    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"Error: file not found: {md_path}")
        sys.exit(1)

    out_path = Path(sys.argv[2]) if len(sys.argv) >= 3 else md_path.with_suffix(".docx")
    result = Converter().convert(md_path, out_path)
    print(f"DOCX written to {result} ({result.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
